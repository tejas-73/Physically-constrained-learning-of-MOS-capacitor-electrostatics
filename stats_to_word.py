import docx
import os


def create_document(model_name, path_list):
    if not os.path.exists(f'./{model_name}_outputs.docx'):
        document = docx.Document()
    else:
        document = docx.Document(f'./{model_name}_outputs.docx')
    for img_path in path_list:
        document.add_picture(f'{img_path}')
    document.save(f'{model_name}_outputs.docx')


